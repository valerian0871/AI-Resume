from io import BytesIO
from docx import Document


def build_resume_docx(optimized_resume_draft: dict) -> BytesIO:
    document = Document()

    document.add_heading("Optimized Resume Draft", level=0)

    professional_summary = optimized_resume_draft.get("professional_summary", "")
    key_skills = optimized_resume_draft.get("key_skills", [])
    experience_bullets = optimized_resume_draft.get("experience_bullets", [])
    project_bullets = optimized_resume_draft.get("project_bullets", [])

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(professional_summary)

    document.add_heading("Key Skills", level=1)
    for skill in key_skills:
        document.add_paragraph(skill, style="List Bullet")

    document.add_heading("Experience", level=1)
    for bullet in experience_bullets:
        document.add_paragraph(bullet, style="List Bullet")

    document.add_heading("Projects", level=1)
    for bullet in project_bullets:
        document.add_paragraph(bullet, style="List Bullet")

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return file_stream